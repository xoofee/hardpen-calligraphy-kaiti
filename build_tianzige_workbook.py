import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
IMG_DIR = OUT_DIR / "tianzige_rows"
DOCX_PATH = OUT_DIR / "tianzige_full_workbook.docx"
CHAR_CACHE = OUT_DIR / "common_chars_level_1.txt"

FIRST_PAGE_CHARS = ["一", "二", "三", "十", "人", "大", "小", "口"]
SOURCE_URLS = [
    "https://www.hanyuguoxue.com/zidian/guifanhanzi-sn-1",
    "https://www.hanyuguoxue.com/zidian/guifanhanzi-sn-1-p2",
    "https://www.hanyuguoxue.com/zidian/guifanhanzi-sn-1-p3",
    "https://www.hanyuguoxue.com/zidian/guifanhanzi-sn-1-p4",
]

CELL_PX = 420
CELL_INCH = 0.55
COLS = 12
ROWS_PER_PAGE = 16
ROW_GAP_INCH = 0.07
ROW_WIDTH_INCH = CELL_INCH * COLS

FONT_CANDIDATES = [
    "/System/Library/Fonts/STKaiti.ttc",
    "/System/Library/Fonts/Supplemental/STKaiti.ttc",
    "/System/Library/Fonts/Supplemental/Kaiti.ttc",
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/88d6cc32a907955efa1d014207889413890573be.asset/AssetData/Kaiti.ttc",
    "/System/Library/AssetsV2/PreinstalledAssetsV2/InstallWithOs/com_apple_MobileAsset_Font7/11a76f9d5fd800227415e64d90ddc450c8acac3e.asset/AssetData/Kaiti.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/PingFang.ttc",
]


def find_font():
    for candidate in FONT_CANDIDATES:
        if Path(candidate).exists():
            return candidate
    return None


def font_for_char(font_path, size):
    if font_path:
        try:
            return ImageFont.truetype(font_path, size=size, index=0)
        except Exception:
            pass
    return ImageFont.load_default()


def fit_font_for_char(draw, font_path, char, target_px):
    size = int(target_px)
    while size < CELL_PX:
        font = font_for_char(font_path, size)
        bbox = draw.textbbox((0, 0), char, font=font)
        if max(bbox[2] - bbox[0], bbox[3] - bbox[1]) >= target_px:
            return font
        size += 6
    return font_for_char(font_path, int(target_px))


def draw_cell(draw, x0, y0, char=None, color=(28, 28, 28), font_path=None):
    border = (46, 46, 46)
    guide = (178, 178, 178)
    pad = 7
    mid_x = x0 + CELL_PX // 2
    mid_y = y0 + CELL_PX // 2
    x1 = x0 + CELL_PX
    y1 = y0 + CELL_PX

    draw.rectangle([x0 + pad, y0 + pad, x1 - pad, y1 - pad], outline=border, width=4)
    draw.line([mid_x, y0 + pad, mid_x, y1 - pad], fill=guide, width=2)
    draw.line([x0 + pad, mid_y, x1 - pad, mid_y], fill=guide, width=2)

    if char:
        font = fit_font_for_char(draw, font_path, char, int(CELL_PX * 0.80))
        bbox = draw.textbbox((0, 0), char, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        tx = x0 + (CELL_PX - tw) / 2 - bbox[0]
        ty = y0 + (CELL_PX - th) / 2 - bbox[1] - 4
        draw.text((tx, ty), char, fill=color, font=font)


def draw_row(char, font_path):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / f"row_{ord(char)}.png"
    if path.exists():
        return path

    width = CELL_PX * COLS
    height = CELL_PX
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    cell_chars = [char] + [char] * 5 + [None] * 6
    colors = [(25, 25, 25)] + [(196, 196, 196)] * 5 + [(25, 25, 25)] * 6
    for col, value in enumerate(cell_chars):
        draw_cell(draw, col * CELL_PX, 0, value, colors[col], font_path)
    img.save(path, dpi=(300, 300), optimize=True)
    return path


def fetch_common_chars():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if CHAR_CACHE.exists():
        cached = CHAR_CACHE.read_text(encoding="utf-8").strip()
        if cached:
            return list(cached)

    chars = []
    seen = set()
    pattern = re.compile(r'<a title="[^"]+" class=han href=[^>]+>\s*<span>\s*([^<\s])\s*</span>')
    for url in SOURCE_URLS:
        with urllib.request.urlopen(url, timeout=30) as response:
            html = response.read().decode("utf-8")
        for char in pattern.findall(html):
            if char not in seen:
                seen.add(char)
                chars.append(char)

    if len(chars) < 3000:
        raise RuntimeError(f"Expected thousands of common characters, got {len(chars)}")

    CHAR_CACHE.write_text("".join(chars), encoding="utf-8")
    return chars


def workbook_chars():
    official = fetch_common_chars()
    ordered = []
    seen = set()
    for char in FIRST_PAGE_CHARS + official:
        if char not in seen:
            ordered.append(char)
            seen.add(char)
    return ordered


def add_row_image(doc, row_path):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(ROW_GAP_INCH * 72)
    run = paragraph.add_run()
    run.add_picture(str(row_path), width=Inches(ROW_WIDTH_INCH))


def set_cell_border_none(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, margin=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def add_page_table(doc, row_paths):
    table = doc.add_table(rows=len(row_paths), cols=1)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row, row_path in zip(table.rows, row_paths):
        row.height = Inches(CELL_INCH + ROW_GAP_INCH)
        cell = row.cells[0]
        cell.width = Inches(ROW_WIDTH_INCH)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border_none(cell)
        set_cell_margins(cell, 0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(str(row_path), width=Inches(ROW_WIDTH_INCH))


def apply_page_setup(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)


def make_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    font_path = find_font()
    chars = workbook_chars()

    doc = Document()
    apply_page_setup(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Kaiti SC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Kaiti SC")
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(0)

    pages = [chars[i : i + ROWS_PER_PAGE] for i in range(0, len(chars), ROWS_PER_PAGE)]
    for page_index, page_chars in enumerate(pages):
        if page_index:
            apply_page_setup(doc.add_section(WD_SECTION.NEW_PAGE))
        add_page_table(doc, [draw_row(char, font_path) for char in page_chars])

    doc.save(DOCX_PATH)
    return DOCX_PATH, len(chars)


if __name__ == "__main__":
    path, count = make_docx()
    print(path)
    print(f"characters={count}")
