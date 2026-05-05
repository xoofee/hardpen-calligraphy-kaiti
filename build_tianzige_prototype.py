from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
IMG_DIR = OUT_DIR / "tianzige_cells"
DOCX_PATH = OUT_DIR / "tianzige_first_page.docx"

CHARS = ["一", "二", "三", "十", "人", "大", "小", "口", "乙", "丁", "厂", "七", "卜", "八", "入", "儿"]
CELL_PX = 420
CELL_INCH = 0.55

FONT_CANDIDATES = [
    "/System/Library/Fonts/STKaiti.ttc",
    "/System/Library/Fonts/Supplemental/STKaiti.ttc",
    "/System/Library/Fonts/Supplemental/Kaiti.ttc",
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/88d6cc32a907955efa1d014207889413890573be.asset/AssetData/Kaiti.ttc",
    "/System/Library/AssetsV2/PreinstalledAssetsV2/InstallWithOs/com_apple_MobileAsset_Font7/11a76f9d5fd800227415e64d90ddc450c8acac3e.asset/AssetData/Kaiti.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/PingFang.ttc",
]


def find_font():
    for candidate in FONT_CANDIDATES:
        path = Path(candidate)
        if path.exists():
            return str(path)
    return None


def font_for_char(font_path, size):
    if font_path:
        try:
            return ImageFont.truetype(font_path, size=size, index=0)
        except Exception:
            pass
    return ImageFont.load_default()


def fit_font_for_char(draw, font_path, char, target_px):
    size = int(target_px)
    while size < CELL_PX:
        font = font_for_char(font_path, size)
        bbox = draw.textbbox((0, 0), char, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        if max(tw, th) >= target_px:
            return font
        size += 6
    return font_for_char(font_path, int(target_px))


def draw_tianzige(char=None, color=(28, 28, 28), suffix="blank"):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (CELL_PX, CELL_PX), "white")
    draw = ImageDraw.Draw(img)

    border = (46, 46, 46)
    guide = (178, 178, 178)
    w = CELL_PX
    pad = 7
    mid = w // 2

    draw.rectangle([pad, pad, w - pad, w - pad], outline=border, width=4)
    draw.line([mid, pad, mid, w - pad], fill=guide, width=2)
    draw.line([pad, mid, w - pad, mid], fill=guide, width=2)

    if char:
        font = fit_font_for_char(draw, find_font(), char, int(CELL_PX * 0.80))
        bbox = draw.textbbox((0, 0), char, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        x = (w - tw) / 2 - bbox[0]
        y = (w - th) / 2 - bbox[1] - 4
        draw.text((x, y), char, fill=color, font=font)

    name = f"{ord(char) if char else 'blank'}_{suffix}.png"
    path = IMG_DIR / name
    img.save(path, dpi=(300, 300))
    return path


def set_cell_border_none(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, margin=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def make_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    model_imgs = {c: draw_tianzige(c, (25, 25, 25), "model") for c in CHARS}
    trace_imgs = {c: draw_tianzige(c, (196, 196, 196), "trace") for c in CHARS}
    blank_img = draw_tianzige(None)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)

    styles = doc.styles
    styles["Normal"].font.name = "Kaiti SC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Kaiti SC")

    table = doc.add_table(rows=len(CHARS), cols=12)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, char in enumerate(CHARS):
        row = table.rows[row_idx]
        row.height = Inches(CELL_INCH + 0.02)
        images = [model_imgs[char]] + [trace_imgs[char]] * 5 + [blank_img] * 6
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(CELL_INCH)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border_none(cell)
            set_cell_margins(cell, 28)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run()
            run.add_picture(str(images[col_idx]), width=Inches(CELL_INCH))

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(make_docx())
